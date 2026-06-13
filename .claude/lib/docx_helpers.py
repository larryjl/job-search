"""
Shared python-docx helpers for all document-generating skills.

This is the single executable source of truth for fonts, colours, margins,
and helper functions. Skills import from here — do not copy-paste these
definitions inline.

Usage:
    import sys, os
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../.claude/lib'))
    from docx_helpers import (
        NAVY, BLACK, FONT,
        set_spacing, add_run, set_margins, keep_with_next,
        add_company_date_row, add_rule, add_bullet,
        add_prof_dev_item, set_metadata,
    )
"""

import os
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

NAVY  = RGBColor(0x1F, 0x37, 0x65)   # #1F3765 — headings, name, accent
BLACK = RGBColor(0x1A, 0x1A, 0x1A)   # #1A1A1A — body text, dates, contact
FONT  = "Arial"


# ---------------------------------------------------------------------------
# Core helpers
# ---------------------------------------------------------------------------

def set_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph space_before, space_after, and optionally line_spacing (all in pt)."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def add_run(paragraph, text, bold=False, size=11, color=None):
    """
    Add a run to paragraph with standard font settings.

    Args:
        paragraph: docx Paragraph object
        text (str): Run text
        bold (bool): Bold weight
        size (int|float): Font size in pt
        color: RGBColor — defaults to BLACK

    Returns:
        Run object
    """
    if color is None:
        color = BLACK
    run = paragraph.add_run(text)
    run.bold           = bold
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.font.name      = FONT
    return run


def set_margins(doc, inches=1):
    """Set equal margins on all sides (default 1 inch)."""
    for section in doc.sections:
        section.top_margin    = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin   = Inches(inches)
        section.right_margin  = Inches(inches)


def keep_with_next(paragraph):
    """Prevent a page break between this paragraph and the next."""
    paragraph.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# Resume-specific helpers
# ---------------------------------------------------------------------------

def add_company_date_row(doc, company, date_range):
    """
    Add a paragraph with company name left-aligned and date range flush right.

    Uses a right tab stop at 6.5 inches (standard letter width minus margins).

    Args:
        doc: Document object
        company (str): Company name (bold, BLACK)
        date_range (str): Date range string (regular, BLACK)

    Returns:
        Paragraph object
    """
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)

    # Add right tab stop at 6.5 inches
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(6.5 * 1440)))  # twips
    tabs.append(tab)
    pPr.append(tabs)

    run_co = p.add_run(company)
    run_co.bold           = True
    run_co.font.color.rgb = BLACK
    run_co.font.name      = FONT

    p.add_run('\t')

    run_date = p.add_run(date_range)
    run_date.bold           = False
    run_date.font.color.rgb = BLACK
    run_date.font.name      = FONT

    return p


def add_prof_dev_item(doc, credential, description=None):
    """
    Add a Professional Development entry as two separate paragraphs.

    The credential is bold; the description (if any) is a plain paragraph
    on the next line. Never concatenate them inline.

    Args:
        doc: Document object
        credential (str): e.g. "dbt Analytics Engineering | 2024"
        description (str|None): Optional plain-text description line

    Returns:
        tuple: (credential_paragraph, description_paragraph or None)
    """
    p_cred = doc.add_paragraph()
    set_spacing(p_cred, before=0, after=0 if description else 6)
    add_run(p_cred, credential, bold=True)

    p_desc = None
    if description:
        p_desc = doc.add_paragraph()
        set_spacing(p_desc, before=0, after=6)
        add_run(p_desc, description)

    return p_cred, p_desc


# ---------------------------------------------------------------------------
# Requirements-matrix helpers
# ---------------------------------------------------------------------------

def add_rule(doc, color="2E5496", size=12):
    """
    Draw a horizontal rule using a bottom border on a blank paragraph.

    Args:
        doc: Document object
        color (str): Hex color string without '#' (default: "2E5496")
        size (int): Border size in eighths of a point (default: 12)

    Returns:
        Paragraph object
    """
    p   = doc.add_paragraph()
    set_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, indent=0.25):
    """
    Add a hanging-indent bullet point paragraph.

    Args:
        doc: Document object
        text (str): Bullet text (bullet char prepended automatically)
        indent (float): Left indent in inches (default 0.25)

    Returns:
        Paragraph object
    """
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2)
    p.paragraph_format.left_indent       = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-indent)
    add_run(p, "• " + text)
    return p


# ---------------------------------------------------------------------------
# Metadata
# ---------------------------------------------------------------------------

def set_metadata(doc, project_root=None):
    """
    Set standard docx core properties from profile/master-resume.md.

    Reads the first line of master-resume.md (strips leading '# ') to get
    the candidate name, then sets author, last_modified_by, and clears comments.

    Args:
        doc: Document object
        project_root (str|None): Override project root path. If None, resolved
            via project_paths.get_project_root().

    Returns:
        str: Candidate name that was set
    """
    if project_root is None:
        try:
            from .project_paths import get_project_root
        except ImportError:
            from project_paths import get_project_root
        project_root = get_project_root()

    resume_path = os.path.join(project_root, "profile", "master-resume.md")
    with open(resume_path) as f:
        candidate_name = f.readline().lstrip("# ").strip() or "Unknown"

    doc.core_properties.author           = candidate_name
    doc.core_properties.last_modified_by = candidate_name
    doc.core_properties.comments         = ""
    return candidate_name
