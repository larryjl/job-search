#!/usr/bin/env python3
"""
Generate a tailored resume .docx (and .pdf) from a structured content dict.

Overview
--------
The script follows a build → validate → trim loop:

1. **Build** — ``build_resume()`` populates a fresh ``python-docx`` Document
   object from the structured ``content`` dict (summary, skills, experience,
   education, professional development, optional sections).
2. **Validate** — the DOCX is saved and immediately re-opened to verify it is
   a well-formed file that Word/LibreOffice can read without errors.
3. **Convert & measure** — ``soffice`` converts the DOCX to PDF and
   ``pdfinfo`` counts the pages.
4. **Trim if needed** — if the PDF exceeds ``--max-pages`` (default 2),
   ``trim_one_bullet()`` removes the weakest bullet from the oldest role and
   the loop repeats from step 1 with the mutated content dict.  If no bullets
   remain to trim, ``trim_summary()`` shortens the summary.  The loop runs up
   to ``max_passes`` times before accepting whatever page count results.

Usage:
    python3 skills/tailor-resume/generate_resume.py \\
        --company "telus" \\
        --role "data-strategist" \\
        --content-file /tmp/resume_content.py \\
        --output-dir job-outputs/resumes/

Content file format (/tmp/resume_content.py):
    content = {
        "summary": "...",
        "skills": [
            {"category": "Analytics Engineering", "items": ["dbt", "Snowflake"]},
        ],
        "experience": [
            {
                "title": "Senior Data Analyst",
                "company": "ACME Corp",
                "location": "Calgary, AB",
                "dates": "Jan 2022 – Present",
                "bullets": ["Led...", "Built..."],
            }
        ],
        "education": [
            {
                "degree": "Bachelor of Commerce",
                "institution": "University of Calgary",
                "location": "Calgary, AB",
                "dates": "2012 – 2016",
                "details": [],   # optional: list of strings
            }
        ],
        "prof_dev": [
            {"credential": "dbt Analytics Engineering | 2024", "description": "Data modeling"},
            {"credential": "AWS Solutions Architect | 2023", "description": None},
        ],
        "optional_sections": [],  # list of {"heading": "...", "items": [...str]}
    }

Output:
    Prints SAVED:/path/to/file.docx and SAVED:/path/to/file.pdf on success.
    Prints ERROR:... and exits with code 1 on failure.
    Prints PAGE_COUNT:N after PDF conversion.
"""

import argparse
import importlib.util
import os
import subprocess
import sys

# ---------------------------------------------------------------------------
# Path setup — works when run from any cwd
# ---------------------------------------------------------------------------

# Resolve the directory containing this script so relative lib references
# are stable regardless of the working directory the caller uses.
SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = None  # resolved lazily via get_project_root()


def _setup_lib_path() -> None:
    """Prepend the shared lib directory to sys.path if not already present."""
    lib_path = os.path.normpath(os.path.join(SCRIPT_DIR, '../lib'))
    if lib_path not in sys.path:
        sys.path.insert(0, lib_path)


_setup_lib_path()

from project_paths import get_project_root  # noqa: E402
from docx_helpers import (                  # noqa: E402
    NAVY, BLACK, FONT, FONT_NAME,
    set_spacing, add_run, set_margins, keep_with_next,
    add_section_heading, add_company_date_row, add_prof_dev_item, set_metadata,
)
from docx import Document                   # noqa: E402
from docx.shared import Pt                  # noqa: E402


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_content(content_file: str) -> dict:
    """
    Dynamically load and return the ``content`` dict from a Python source file.

    The content file is a plain ``.py`` file (not a JSON or YAML file) so it
    can contain Python expressions like list comprehensions if needed.
    ``importlib.util`` is used instead of ``exec()`` or ``import`` because it
    lets us load an arbitrary file path as a module without adding its
    directory to ``sys.path`` or risking name collisions with real modules.

    Args:
        content_file: Absolute path to the Python file that defines ``content``.

    Returns:
        The ``content`` dict defined in that file.

    Raises:
        ValueError: If the file does not define a variable named ``content``.
    """
    # ``spec_from_file_location`` creates a module spec that tells Python how
    # to load the file — essentially "treat this file as a module called
    # 'resume_content'".
    spec = importlib.util.spec_from_file_location("resume_content", content_file)

    # ``module_from_spec`` allocates an empty module object from that spec.
    module = importlib.util.module_from_spec(spec)

    # ``exec_module`` actually runs the file, populating ``module`` with all
    # variables defined at the top level (including ``content``).
    spec.loader.exec_module(module)

    if not hasattr(module, 'content'):
        raise ValueError(
            f"content_file must define a variable named 'content': {content_file}"
        )
    return module.content


def get_pdf_page_count(pdf_path: str) -> int | None:
    """
    Return the page count of a PDF file using the ``pdfinfo`` CLI tool.

    Args:
        pdf_path: Absolute path to the PDF file.

    Returns:
        Integer page count, or ``None`` if ``pdfinfo`` is unavailable or fails.
    """
    try:
        result = subprocess.run(
            ["pdfinfo", pdf_path],
            capture_output=True, text=True, timeout=15,
        )
        for line in result.stdout.splitlines():
            if line.startswith("Pages:"):
                return int(line.split(":")[1].strip())
    except Exception:
        pass
    return None


def convert_to_pdf(docx_path: str, output_dir: str) -> str | None:
    """
    Convert a DOCX file to PDF using LibreOffice (``soffice``).

    soffice creates temporary lock files and a ``.~lock`` file in the same
    directory as the output PDF while the conversion is running.  If we
    converted directly into ``job-outputs/resumes/`` those artefacts would
    clutter the resumes folder and could confuse file-listing tools.  Instead
    we convert into a fresh temporary directory and then move only the finished
    PDF to the target output directory.

    Args:
        docx_path:  Absolute path to the source DOCX file.
        output_dir: Directory where the final PDF should land.

    Returns:
        Absolute path to the saved PDF, or ``None`` if conversion failed.
    """
    import shutil
    import tempfile

    # Create an isolated scratch directory for soffice's temporary artefacts.
    tmp_dir = tempfile.mkdtemp(prefix="resume_pdf_")
    try:
        result = subprocess.run(
            [
                "soffice", "--headless",
                "--convert-to", "pdf",
                docx_path,
                "--outdir", tmp_dir,   # write PDF + lock files here, not to output_dir
            ],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            return None

        # soffice names the PDF after the DOCX basename with the extension swapped.
        tmp_pdf = os.path.join(
            tmp_dir,
            os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
        )
        if not os.path.exists(tmp_pdf):
            return None

        # Move the clean PDF to the intended output directory.
        dest_pdf = os.path.join(output_dir, os.path.basename(tmp_pdf))
        shutil.move(tmp_pdf, dest_pdf)
        return dest_pdf
    finally:
        # Always clean up the temp dir — including any soffice lock files —
        # even if an exception was raised during conversion.
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_resume(doc: Document, content: dict, project_root: str) -> None:
    """
    Populate a blank ``python-docx`` Document with resume content.

    Renders sections in this order: header (name + contact), summary, skills,
    professional experience, education, professional development, optional
    sections.  Sections absent from ``content`` are silently skipped.

    Args:
        doc:          A freshly created ``Document`` object with margins already set.
        content:      Structured content dict (see module docstring for schema).
        project_root: Absolute path to the job-search project root, used by
                      ``set_metadata`` to read the candidate's name from the
                      master resume.
    """
    # Write author/title into the DOCX core properties and return the
    # candidate's name so we can use it in the document header.
    candidate_name = set_metadata(doc, project_root)

    # --- Header ---
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=4)

    # The candidate's name is the only element rendered in Georgia (FONT_NAME).
    # All other text uses the default body font (Calibri).  The Georgia serif
    # gives the name a more distinctive, professional feel at large size.
    run_name = add_run(p, candidate_name, bold=True, size=16, color=NAVY)
    run_name.font.name = FONT_NAME  # Georgia — name only; rest of doc uses default font

    # ``contact_line`` is optional: skills that call this script may include it
    # (e.g. "Calgary, AB | leelawrencej@gmail.com | linkedin.com/in/...").
    # If absent, the header is just the name — no blank line is added.
    contact_line = content.get("contact_line", "")
    if contact_line:
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=12)
        add_run(p, contact_line, size=10)

    # --- Summary ---
    summary = content.get("summary", "")
    if summary:
        add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=6)
        add_run(p, summary)

    # --- Skills ---
    skills = content.get("skills", [])
    if skills:
        add_section_heading(doc, "Skills")
        for skill_group in skills:
            p = doc.add_paragraph()
            set_spacing(p, before=0, after=2)
            # Bold category label (e.g. "Analytics Engineering: ") followed by
            # comma-separated items in normal weight.
            add_run(p, skill_group["category"] + ": ", bold=True)
            add_run(p, ", ".join(skill_group["items"]))

    # --- Experience ---
    experience = content.get("experience", [])
    if experience:
        add_section_heading(doc, "Professional Experience")
        for job in experience:
            # Job title — bold, with extra space above to separate from the
            # previous role's last bullet.
            p = doc.add_paragraph()
            set_spacing(p, before=6, after=0)
            # ``keep_with_next`` tells Word/LibreOffice to never insert a page
            # break between this paragraph and the one that follows it.  Applied
            # to the title so the company/date row is never stranded on the
            # next page without its heading.
            keep_with_next(p)
            add_run(p, job["title"], bold=True)

            # Company + location on the left, date range on the right.
            # ``add_company_date_row`` uses a tab stop at the right margin so
            # the date is right-aligned without needing a table.
            company_loc = job["company"]
            if job.get("location"):
                company_loc += f", {job['location']}"
            row = add_company_date_row(doc, company_loc, job.get("dates", ""))
            # Keep the company row with the first bullet so the header block
            # (title + company/date) never becomes orphaned at the bottom of a page.
            keep_with_next(row)

            # Bullets — hanging indent: left edge at 18pt (~0.25in), first line
            # pulled back 18pt so the bullet character sits at the left margin.
            for bullet in job.get("bullets", []):
                p = doc.add_paragraph()
                set_spacing(p, before=0, after=2)
                p.paragraph_format.left_indent       = Pt(18)   # body text indented 0.25in
                p.paragraph_format.first_line_indent = Pt(-18)  # bullet hangs back to margin
                add_run(p, "•  " + bullet)

    # --- Education ---
    education = content.get("education", [])
    if education:
        add_section_heading(doc, "Education")
        for edu in education:
            # Institution + location on the left, dates on the right (same
            # right-aligned tab-stop pattern as experience rows).
            institution_loc = edu["institution"]
            if edu.get("location"):
                institution_loc += f", {edu['location']}"
            row = add_company_date_row(doc, institution_loc, edu.get("dates", ""))
            keep_with_next(row)

            # Degree on its own line, bold.
            p = doc.add_paragraph()
            set_spacing(p, before=0, after=2)
            add_run(p, edu["degree"], bold=True)

            # Optional detail lines (honours, GPA, relevant coursework).
            for detail in edu.get("details", []):
                p2 = doc.add_paragraph()
                set_spacing(p2, before=0, after=2)
                add_run(p2, "•  " + detail)

    # --- Professional Development ---
    prof_dev = content.get("prof_dev", [])
    if prof_dev:
        add_section_heading(doc, "Professional Development")
        for item in prof_dev:
            add_prof_dev_item(doc, item["credential"], item.get("description"))

    # --- Optional sections (certifications, volunteer work, publications, etc.) ---
    for section in content.get("optional_sections", []):
        add_section_heading(doc, section["heading"])
        for item in section.get("items", []):
            p = doc.add_paragraph()
            set_spacing(p, before=0, after=2)
            add_run(p, item)


def trim_one_bullet(content: dict) -> bool:
    """
    Drop the last bullet from the oldest role that still has more than one bullet.

    Strategy — oldest role first:
    Older roles are less relevant to recent job applications, so their bullets
    are the cheapest to sacrifice.  We start from the end of the experience
    list (oldest) and work backwards toward the present.  The most recent role
    (index 0) is protected and only trimmed as a last resort.

    "Weakest bullet" is defined as the *last* bullet in the list.  The caller
    is expected to have ordered bullets from most to least impactful, so the
    last one is the weakest.

    Never removes the only bullet in a role (a role with zero bullets would be
    malformed).

    Args:
        content: The live content dict (mutated in place if a bullet is removed).

    Returns:
        True if a bullet was removed, False if no roles have a spare bullet to trim.
    """
    experience = content.get("experience", [])
    if not experience:
        return False

    # Iterate from the oldest role (last in list) toward the second-most-recent,
    # skipping index 0 (most recent role) — it is protected in this pass.
    for job in reversed(experience[1:]):
        if len(job.get("bullets", [])) > 1:
            job["bullets"].pop()  # remove the last (weakest) bullet in place
            return True

    # Every older role is already down to 1 bullet.  As a last resort, trim
    # from the most recent role rather than accepting an oversized resume.
    if len(experience[0].get("bullets", [])) > 1:
        experience[0]["bullets"].pop()
        return True

    # All roles are at exactly 1 bullet — nothing safe to remove.
    return False


def trim_summary(content: dict) -> bool:
    """
    Shorten the professional summary by removing its last sentence.

    Args:
        content: The live content dict (mutated in place if a sentence is removed).

    Returns:
        True if a sentence was removed, False if the summary has only one sentence.
    """
    summary = content.get("summary", "")
    sentences = [s.strip() for s in summary.split(".") if s.strip()]
    if len(sentences) > 1:
        content["summary"] = ". ".join(sentences[:-1]) + "."
        return True
    return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    """
    Entry point: parse args, load content, run the build→validate→trim loop.

    The loop runs up to ``max_passes`` times.  Each iteration:
      1. Builds a fresh Document from the (potentially already trimmed) content dict.
      2. Saves and validates the DOCX.
      3. Converts to PDF and counts pages.
      4. If within the page limit — breaks and reports success.
      5. If over the limit — trims one bullet (or one summary sentence) and loops.

    The content dict is mutated by trim calls, so each rebuild reflects the
    cumulative effect of all trims applied so far.
    """
    parser = argparse.ArgumentParser(description="Generate tailored resume docx + pdf")
    parser.add_argument("--company",      required=True,  help="Normalized company name (e.g. 'telus')")
    parser.add_argument("--role",         required=True,  help="Normalized role slug (e.g. 'data-strategist')")
    parser.add_argument("--content-file", required=True,  help="Path to Python file defining 'content' dict")
    parser.add_argument("--output-dir",   required=False, help="Output directory (default: job-outputs/resumes/)")
    parser.add_argument("--max-pages",    type=int, default=2, help="Maximum page count (default: 2)")
    args = parser.parse_args()

    try:
        project_root = get_project_root()
    except RuntimeError as e:
        print(f"ERROR:{e}", file=sys.stderr)
        sys.exit(1)

    output_dir = args.output_dir or os.path.join(project_root, "job-outputs", "resumes")
    os.makedirs(output_dir, exist_ok=True)

    from datetime import date
    today     = date.today().strftime("%Y-%m-%d")
    filename  = f"resume_{args.company}_{args.role}_{today}"
    docx_path = os.path.join(output_dir, filename + ".docx")
    pdf_path  = os.path.join(output_dir, filename + ".pdf")

    try:
        content = load_content(args.content_file)
    except Exception as e:
        print(f"ERROR:Failed to load content file: {e}", file=sys.stderr)
        sys.exit(1)

    # --- Build → validate → trim loop ---
    # We cap iterations at max_passes to guarantee termination even if the
    # resume content is too long to fit in max_pages no matter how much we trim.
    max_passes = 10
    for attempt in range(max_passes):

        # Each iteration builds a completely fresh Document so that trimmed
        # content from previous passes doesn't leave ghost paragraphs.
        doc = Document()
        set_margins(doc)

        try:
            build_resume(doc, content, project_root)
        except Exception as e:
            print(f"ERROR:Document build failed: {e}", file=sys.stderr)
            sys.exit(1)

        # Save the DOCX, then immediately re-open it as a basic sanity check.
        # python-docx can occasionally produce invalid XML for certain content
        # combinations; the re-open catches that before we waste time on PDF
        # conversion.
        doc.save(docx_path)
        try:
            Document(docx_path)
        except Exception as e:
            print(f"ERROR:Generated docx failed validation: {e}", file=sys.stderr)
            sys.exit(1)

        # Convert to PDF via LibreOffice (soffice).
        result_pdf = convert_to_pdf(docx_path, output_dir)
        if not result_pdf:
            print("ERROR:PDF conversion failed (soffice). docx saved but no pdf.", file=sys.stderr)
            print(f"SAVED:{docx_path}")
            sys.exit(1)

        # Report page count to the caller so the agent can log it.
        pages = get_pdf_page_count(result_pdf)
        if pages is not None:
            print(f"PAGE_COUNT:{pages}")
        else:
            # pdfinfo unavailable — we can't measure, so we accept the result.
            pages = 0

        # Within page limit (or page count unknown) — done.
        if pages <= args.max_pages or pages == 0:
            break

        # Over the page limit: trim one bullet from the oldest eligible role.
        # If no bullets can be trimmed, fall back to shortening the summary.
        # If neither trim is possible, accept the oversized resume and stop.
        print(f"INFO:Resume is {pages} pages — trimming (pass {attempt + 1})", file=sys.stderr)
        trimmed = trim_one_bullet(content)
        if not trimmed:
            trimmed = trim_summary(content)
        if not trimmed:
            print(f"INFO:Nothing left to trim — accepting {pages} pages", file=sys.stderr)
            break
        # Loop continues: next iteration rebuilds the document with updated content.

    # Always report the DOCX path; PDF path is only reported if it exists.
    print(f"SAVED:{docx_path}")
    if os.path.exists(pdf_path):
        print(f"SAVED:{pdf_path}")


if __name__ == "__main__":
    main()
