#!/usr/bin/env python3
"""
Generate interview prep notes .docx from a structured content dict.

Usage:
    python3 skills/interview-prep/generate_interview_notes.py \
        --company "telus" \
        --role "data-strategist" \
        --content-file /tmp/interview_content.py \
        --output-dir job-outputs/interview-notes/

Content file format (/tmp/interview_content.py):
    content = {
        "company": "Telus",             # display name
        "role": "Data Strategist",      # display name
        "questions": [
            {
                "number": 1,
                "type": "Behavioural",
                "question": "Tell me about a time you led a complex data project.",
                "approach": "Use the Snowflake migration story. Lead with scope...",
                "story": {
                    "situation": "...",
                    "task": "...",
                    "action": "...",
                    "result": "...",
                    "reflection": "...",
                },
                "tip": "Keep the technical detail brief — they want leadership signals.",
            },
        ],
        "red_flag_questions": [
            {
                "question": "You've been independent for a few years — how do you work within a structure?",
                "why_asked": "Culture fit, coachability",
                "how_to_answer": "Reframe operator independence as a scaling asset...",
            },
        ],
        "company_research": [
            "Telus Health acquired LifeWorks in 2022 — now largest health benefits platform in Canada.",
            "Current focus: AI-driven health outcomes, reducing claims processing time.",
        ],
        "questions_to_ask": [
            "What does success look like for this role in the first 90 days?",
            "How does the data team interface with product and engineering?",
        ],
    }

Output:
    Prints SAVED:/path/to/file.docx on success.
    Prints ERROR:... and exits with code 1 on failure.
"""

import argparse
import importlib.util
import os
import sys
from datetime import date

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def _setup_lib_path() -> None:
    """Add the shared skills/lib directory to sys.path so shared modules are importable."""
    lib_path = os.path.normpath(os.path.join(SCRIPT_DIR, '../lib'))
    if lib_path not in sys.path:
        sys.path.insert(0, lib_path)


_setup_lib_path()

from project_paths import get_project_root   # noqa: E402
from docx_helpers import set_metadata        # noqa: E402
from docx import Document                    # noqa: E402
from docx.shared import Inches, Pt, RGBColor # noqa: E402
from docx.oxml.ns import qn                  # noqa: E402
from docx.oxml import OxmlElement            # noqa: E402

# All text in this document uses Arial for consistent cross-platform rendering
FONT = "Arial"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_content(content_file: str) -> dict:
    """
    Dynamically import a Python file and return its ``content`` dict.

    The content file is a plain .py file (not a package) that defines a
    single top-level variable called ``content``.  Using importlib lets the
    calling skill write an arbitrary Python literal — including multi-line
    strings — without having to serialise through JSON.

    Args:
        content_file: Absolute or relative path to the .py content file.

    Returns:
        The ``content`` dict defined inside the file.

    Raises:
        ValueError: If the loaded module does not define a ``content`` variable.
    """
    # Give the module a stable name so repeated imports don't collide in sys.modules
    spec   = importlib.util.spec_from_file_location("interview_content", content_file)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)

    if not hasattr(module, 'content'):
        raise ValueError(f"content_file must define a variable named 'content': {content_file}")

    return module.content


def add_divider(doc: Document, color: str = "CCCCCC", size: int = 6) -> object:
    """
    Insert a thin horizontal rule between document sections.

    The rule is implemented as a paragraph bottom-border (w:pBdr/w:bottom) using
    Word's Open XML format, rather than a drawn shape, so it flows with the text
    and never causes pagination issues.  The paragraph itself is empty — its only
    purpose is to carry the border styling.

    Args:
        doc:   The python-docx Document to append the divider to.
        color: Hex colour string without the leading '#' (default: light grey "CCCCCC").
        size:  Border thickness in eighths of a point (default 6 = 0.75pt).

    Returns:
        The empty paragraph element that holds the border.
    """
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # w:pBdr wraps all paragraph border definitions; w:bottom is the bottom edge
    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')   # solid line style
    bottom.set(qn('w:sz'),    str(size))  # thickness in 1/8 pt units
    bottom.set(qn('w:space'), '1')        # gap between text and border (points)
    bottom.set(qn('w:color'), color)      # hex colour without '#'
    pBdr.append(bottom)
    pPr.append(pBdr)

    return p


def fix_heading_font(heading: object) -> None:
    """
    Override Word's default heading theme font to use Arial on every run.

    Word's built-in Heading styles inherit their font from the document theme
    (usually Calibri Light or Cambria), which overrides any run-level font set
    through python-docx.  Iterating over the heading's runs after they are
    created and explicitly setting ``run.font.name`` bypasses the theme lookup
    and ensures the heading renders in Arial on all platforms.

    Args:
        heading: A python-docx paragraph that was created with a Heading style.
    """
    for run in heading.runs:
        run.font.name = FONT


def add_body_para(
    doc: Document,
    text: str,
    bold_prefix: str | None = None,
    italic: bool = False,
) -> object:
    """
    Add a plain body paragraph, optionally prefixed with a bold label.

    Splitting the prefix into a separate run is necessary because python-docx
    does not support mixed bold/non-bold within a single ``add_run`` call — each
    run has a single set of character properties.  A bold prefix run followed by
    a normal text run produces the visual pattern "Label: body text" where only
    the label is bold.

    Args:
        doc:         The python-docx Document to append the paragraph to.
        text:        The main (non-bold) paragraph text.
        bold_prefix: Optional label rendered in bold before ``text``
                     (e.g. ``"Situation: "``).  Include trailing space/colon.
        italic:      When True, render ``text`` in italics (used for tips).

    Returns:
        The newly added paragraph.

    Example:
        add_body_para(doc, "Led migration...", bold_prefix="Situation: ")
    """
    p = doc.add_paragraph()

    if bold_prefix:
        # Bold prefix is a separate run so it can carry different character formatting
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(11)

    r2 = p.add_run(text)
    r2.italic = italic
    r2.font.name = FONT
    r2.font.size = Pt(11)

    return p


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_interview_notes(doc: Document, content: dict, project_root: str) -> None:
    """
    Populate a blank python-docx Document with structured interview prep content.

    Sections written in order:
        Part 1 — Top Interview Questions (summary table)
        Part 2 — STAR Story Bank (detailed answers with situation/task/action/result/reflection)
        Part 3 — Company Research (bullet list of key facts)
        Part 4 — Red-Flag Questions (table with reframe guidance)
        Part 5 — Questions to Ask (bullet list)
        Footer  — prompt to run /mock-interview

    Args:
        doc:          A freshly created python-docx Document (no margins set yet —
                      this function sets its own margins; see note below).
        content:      Dict loaded from the content file.
        project_root: Absolute project root path, passed through to set_metadata.
    """
    set_metadata(doc, project_root)

    company = content.get("company", "")
    role    = content.get("role", "")
    today   = date.today().strftime("%B %-d, %Y")

    # Set 1-inch margins on all sides.
    # Note: this script sets margins here directly rather than calling
    # set_margins() from docx_helpers because set_margins() uses 0.85-inch
    # margins tuned for dense resume layouts.  Interview notes are more
    # readable with the standard 1-inch margin, so we set them inline.
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # --- Document title and date ---
    title = doc.add_heading(f"Interview Prep — {role} at {company}", level=1)
    fix_heading_font(title)
    doc.add_paragraph(today).runs[0].font.name = FONT

    # ---- Part 1: Question list table ----
    questions = content.get("questions", [])
    if questions:
        add_divider(doc)
        h = doc.add_heading("Part 1 — Top Interview Questions", level=2)
        fix_heading_font(h)

        # Three-column table: question number | type | question text
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Type"
        hdr[2].text = "Question"

        # Bold the header row cells
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.name = FONT

        # Column widths: narrow index | medium type label | wide question text
        # These are set on each data row because Word's table model applies
        # column widths per-cell rather than per-column in the python-docx API.
        for row_data in questions:
            row = table.add_row().cells
            row[0].width = Inches(0.4)
            row[1].width = Inches(1.1)
            row[2].width = Inches(5.0)
            row[0].text = str(row_data.get("number", ""))
            row[1].text = row_data.get("type", "")
            row[2].text = row_data.get("question", "")

            # Apply font to every run in every cell of this data row
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT

    # ---- Part 2: STAR stories ----
    if questions:
        add_divider(doc)
        h = doc.add_heading("Part 2 — STAR Story Bank", level=2)
        fix_heading_font(h)

        for q in questions:
            story = q.get("story")
            if not story:
                continue  # Skip questions that have no STAR story defined

            qh = doc.add_heading(f"Q{q.get('number', '')} — {q.get('question', '')}", level=3)
            fix_heading_font(qh)

            approach = q.get("approach", "")
            if approach:
                add_body_para(doc, approach, bold_prefix="Approach: ")

            # Iterate through the five STAR + Reflection components in a fixed order
            # so they always appear consistently regardless of dict key ordering.
            for label in ("situation", "task", "action", "result", "reflection"):
                text = story.get(label, "")
                if text:
                    # "reflection" gets a star prefix to visually distinguish it
                    # as a meta-commentary on the story rather than a factual component.
                    # All other labels are capitalised normally (e.g. "Situation: ").
                    prefix = "★ Reflection: " if label == "reflection" else f"{label.capitalize()}: "
                    add_body_para(doc, text, bold_prefix=prefix)

            tip = q.get("tip", "")
            if tip:
                # Tips are rendered in italics to visually separate coaching notes
                # from the story content itself
                add_body_para(doc, f"💡 Tip: {tip}", italic=True)

    # ---- Part 3: Company research ----
    research = content.get("company_research", [])
    if research:
        add_divider(doc)
        h = doc.add_heading("Part 3 — Company Research", level=2)
        fix_heading_font(h)

        for item in research:
            # add_paragraph returns the paragraph; accessing .runs[0] is only safe
            # if Word created at least one run for the bullet text.  The conditional
            # assignment (... if ... else None) silently skips font assignment on
            # empty paragraphs rather than raising an IndexError.
            p = doc.add_paragraph(item, style="List Bullet")
            # Only set the font if Word created at least one run (empty paragraphs
            # have no runs and accessing runs[0] would raise an IndexError).
            if p.runs:
                p.runs[0].font.name = FONT

    # ---- Part 4: Red-flag questions ----
    red_flags = content.get("red_flag_questions", [])
    if red_flags:
        add_divider(doc)
        h = doc.add_heading("Part 4 — Red-Flag Questions", level=2)
        fix_heading_font(h)

        # Three-column table: tricky question | why it's asked | how to reframe it
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Question"
        hdr[1].text = "Why It's Asked"
        hdr[2].text = "How to Answer It"

        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.name = FONT

        for rf in red_flags:
            row = table.add_row().cells
            row[0].text = rf.get("question", "")
            row[1].text = rf.get("why_asked", "")
            row[2].text = rf.get("how_to_answer", "")
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT

    # ---- Part 5: Questions to ask ----
    questions_to_ask = content.get("questions_to_ask", [])
    if questions_to_ask:
        add_divider(doc)
        h = doc.add_heading("Part 5 — Questions to Ask", level=2)
        fix_heading_font(h)

        for item in questions_to_ask:
            p = doc.add_paragraph(item, style="List Bullet")
            # Guard against empty paragraphs that have no runs before setting font
            if p.runs:
                p.runs[0].font.name = FONT

    # --- Footer prompt ---
    add_divider(doc)
    p = doc.add_paragraph()
    r = p.add_run(f"Ready to practise live? Run /mock-interview for {role} at {company}.")
    r.italic = True
    r.font.name = FONT
    r.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    """
    Entry point: parse arguments, build the interview notes document, save docx.

    Exit codes:
        0 — docx saved and validated successfully
        1 — any error (content load failure, build failure, validation failure)

    Output on success:
        SAVED:/absolute/path/to/interview_<company>_<role>_<date>.docx
    The SAVED: prefix is a machine-readable protocol parsed by the calling skill.
    """
    parser = argparse.ArgumentParser(description="Generate interview prep notes docx")
    parser.add_argument("--company",      required=True,  help="Normalised company slug (e.g. 'telus')")
    parser.add_argument("--role",         required=True,  help="Normalised role slug (e.g. 'data-strategist')")
    parser.add_argument("--content-file", required=True,  help="Path to a .py file defining a 'content' dict")
    parser.add_argument("--output-dir",   required=False, help="Destination directory (default: job-outputs/interview-notes/)")
    args = parser.parse_args()

    # Resolve the project root so all subsequent paths are absolute
    try:
        project_root = get_project_root()
    except RuntimeError as e:
        print(f"ERROR:{e}", file=sys.stderr)
        sys.exit(1)

    # Use the supplied output dir or fall back to the standard interview-notes folder
    output_dir = args.output_dir or os.path.join(project_root, "job-outputs", "interview-notes")
    os.makedirs(output_dir, exist_ok=True)

    # Build the output filename following the project naming convention
    today     = date.today().strftime("%Y-%m-%d")
    filename  = f"interview_{args.company}_{args.role}_{today}"
    docx_path = os.path.join(output_dir, filename + ".docx")

    # Load the content dict from the caller-supplied Python file
    try:
        content = load_content(args.content_file)
    except Exception as e:
        print(f"ERROR:Failed to load content file: {e}", file=sys.stderr)
        sys.exit(1)

    # Create a blank document; margins are set inside build_interview_notes
    doc = Document()

    try:
        build_interview_notes(doc, content, project_root)
    except Exception as e:
        print(f"ERROR:Document build failed: {e}", file=sys.stderr)
        sys.exit(1)

    doc.save(docx_path)

    # Re-open the saved file to confirm it is a valid docx (catches corruption)
    try:
        Document(docx_path)
    except Exception as e:
        print(f"ERROR:Generated docx failed validation: {e}", file=sys.stderr)
        sys.exit(1)

    # Emit the SAVED: line so the calling skill can parse the output path
    print(f"SAVED:{docx_path}")


if __name__ == "__main__":
    main()
