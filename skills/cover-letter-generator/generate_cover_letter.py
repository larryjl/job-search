#!/usr/bin/env python3
"""
Generate a cover letter .docx (and .pdf) from a structured content dict.

Usage:
    python3 skills/cover-letter-generator/generate_cover_letter.py \
        --company "shopify" \
        --role "senior-analytics-engineer" \
        --content-file /tmp/cover_content.py \
        --output-dir job-outputs/cover-letters/

Content file format (/tmp/cover_content.py):
    content = {
        "re_line": "Senior Analytics Engineer — Shopify",
        "hiring_manager": "Hiring Team",      # or a real name
        "opening": "...",
        "middle": "...",
        "closing": "...",
        # Optional — read from master resume if omitted:
        "contact_line": "Calgary, AB | email | phone | linkedin",
    }

Output:
    Prints SAVED:/path/to/file.docx and SAVED:/path/to/file.pdf on success.
    Prints ERROR:... and exits with code 1 on failure.

The SAVED: prefix is a machine-readable protocol parsed by the calling skill
(cover-letter-generator/SKILL.md) to locate output files without globbing.
"""

import argparse
import importlib.util
import os
import subprocess
import sys
from datetime import date

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def _setup_lib_path() -> None:
    """Add the shared skills/lib directory to sys.path so shared modules are importable."""
    lib_path = os.path.normpath(os.path.join(SCRIPT_DIR, '../lib'))
    if lib_path not in sys.path:
        sys.path.insert(0, lib_path)


_setup_lib_path()

from project_paths import get_project_root   # noqa: E402
from docx_helpers import (                   # noqa: E402
    NAVY, BLACK,
    set_spacing, add_run, set_margins, set_metadata,
)
from docx import Document                    # noqa: E402


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_content(content_file: str) -> dict:
    """
    Dynamically import a Python file and return its ``content`` dict.

    The content file is a plain .py file (not a package) that defines a
    single top-level variable called ``content``.  Using importlib lets the
    calling skill write an arbitrary Python literal — including multi-line
    strings — without having to serialise through JSON.

    Args:
        content_file: Absolute or relative path to the .py content file.

    Returns:
        The ``content`` dict defined inside the file.

    Raises:
        ValueError: If the loaded module does not define a ``content`` variable.
    """
    # Give the module a stable name so repeated imports don't collide in sys.modules
    spec   = importlib.util.spec_from_file_location("cover_content", content_file)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)

    if not hasattr(module, 'content'):
        raise ValueError(f"content_file must define a variable named 'content': {content_file}")

    return module.content


def convert_to_pdf(docx_path: str, output_dir: str) -> str | None:
    """
    Convert a .docx file to PDF using LibreOffice (soffice) in headless mode.

    Returns the destination PDF path on success, or None on failure.

    Why stage through /tmp instead of converting directly into output_dir:
        soffice creates lock files (.~lock.<filename>#) and temporary files
        alongside the output during conversion.  If we point --outdir at the
        cover-letters folder those lock files end up there and can confuse
        subsequent glob operations.  By converting into a throwaway temp
        directory and moving only the final .pdf over, the output folder stays
        clean regardless of whether the conversion succeeds or crashes.

    Args:
        docx_path:  Absolute path to the source .docx file.
        output_dir: Directory where the final .pdf should land.

    Returns:
        Absolute path to the .pdf in output_dir, or None if conversion failed.
    """
    import shutil
    import tempfile

    # Create a short-lived staging directory that soffice can litter freely
    tmp_dir = tempfile.mkdtemp(prefix="cover_pdf_")
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", tmp_dir],
            capture_output=True, text=True, timeout=60,
        )

        # A non-zero return code means soffice reported a conversion error
        if result.returncode != 0:
            return None

        # soffice names the output file by replacing the source extension with .pdf
        tmp_pdf = os.path.join(tmp_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")

        # Guard against the edge case where soffice exits 0 but wrote no file
        if not os.path.exists(tmp_pdf):
            return None

        # Move (not copy) so we don't duplicate a potentially large file
        dest_pdf = os.path.join(output_dir, os.path.basename(tmp_pdf))
        shutil.move(tmp_pdf, dest_pdf)
        return dest_pdf

    finally:
        # Always remove the staging directory, even if an exception was raised
        shutil.rmtree(tmp_dir, ignore_errors=True)


def read_contact_line_from_resume(project_root: str) -> str:
    """
    Extract the candidate's contact line from the master resume markdown file.

    The master resume follows a two-line header convention:
        Line 1 (first non-empty): ``# Full Name``   ← the H1 heading
        Line 2 (second non-empty): contact info     ← city | email | phone | LinkedIn

    We want the second non-empty line because the first is the candidate's name
    (already extracted separately via set_metadata), while the second contains
    the pipe-separated contact details that appear under the name in the letter.

    Returns an empty string if the file is missing or cannot be read.

    Args:
        project_root: Absolute path to the root of the job-search project.

    Returns:
        The raw second non-empty line from the resume, or ``""`` on any error.
    """
    resume_path = os.path.join(project_root, "profile", "master-resume.md")
    try:
        with open(resume_path) as f:
            # Filter out blank lines so heading + contact are always indices 0 and 1
            lines = [l.strip() for l in f.readlines() if l.strip()]

        # lines[0] = "# Candidate Name", lines[1] = contact info
        return lines[1] if len(lines) > 1 else ""
    except Exception:
        # Silently return empty string; the caller will omit the contact block
        return ""


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_cover_letter(doc: Document, content: dict, project_root: str) -> None:
    """
    Populate a blank python-docx Document with formatted cover letter content.

    Writes sections in order: name header → contact line → date → Re: line →
    salutation → body paragraphs → sign-off → name.

    Args:
        doc:          A freshly created python-docx Document (margins already set).
        content:      Dict loaded from the content file; keys are described in the
                      module docstring above.
        project_root: Absolute project root path, used to fall back to the
                      master resume for the contact line.
    """
    # set_metadata writes author/last_modified_by to doc properties and returns
    # the candidate name string (read from the first line of master-resume.md)
    candidate_name = set_metadata(doc, project_root)

    # --- Name header (16pt bold navy) ---
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2)
    add_run(p, candidate_name, bold=True, size=16, color=NAVY)

    # --- Contact line (10pt, below name) ---
    contact_line = content.get("contact_line") or read_contact_line_from_resume(project_root)
    if contact_line:
        # The resume's second line may start with markdown artifacts like "# " or "| "
        # because some resume templates use a blockquote or heading for the contact row.
        # lstrip removes any leading hash, pipe, or space characters before rendering.
        contact_line = contact_line.lstrip("#| ").strip()
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=12)
        add_run(p, contact_line, size=10)

    # --- Date ---
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=8)
    # "%-d" is a Linux/macOS strftime extension that formats the day without a
    # leading zero (e.g. "June 7, 2026" instead of "June 07, 2026").
    add_run(p, date.today().strftime("%B %-d, %Y"))

    # --- Re: line (bold navy, subject of the letter) ---
    re_line = content.get("re_line", "")
    if re_line:
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=12)
        # "Re: " label and the role title are two separate runs so both are bold navy
        add_run(p, "Re: ", bold=True, color=NAVY)
        add_run(p, re_line, bold=True, color=NAVY)

    # --- Salutation ---
    hiring_manager = content.get("hiring_manager", "Hiring Team")
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=8)
    add_run(p, f"Dear {hiring_manager},")

    # --- Body paragraphs ---
    # The three keys (opening / middle / closing) represent the standard three-paragraph
    # cover letter structure.  Keying by semantic name (rather than a numbered list)
    # lets the content file author omit any section without shifting indices.
    for body_key in ("opening", "middle", "closing"):
        body_text = content.get(body_key, "")
        if body_text:
            p = doc.add_paragraph()
            set_spacing(p, before=0, after=10)
            add_run(p, body_text)

    # --- Sign-off ---
    p = doc.add_paragraph()
    set_spacing(p, before=12, after=0)
    add_run(p, "Sincerely,")

    # 24pt of space before the printed name leaves room for a handwritten signature
    # when the letter is printed (roughly one blank line at the default line height).
    p = doc.add_paragraph()
    set_spacing(p, before=24, after=0)
    add_run(p, candidate_name, bold=True)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    """
    Entry point: parse arguments, build the cover letter, save docx + pdf.

    Exit codes:
        0 — both docx and pdf saved successfully
        1 — any error (content load failure, build failure, PDF conversion failure)

    Output lines on stdout use the ``SAVED:`` prefix so the calling skill can
    parse the output paths without needing to reconstruct the filename logic.
    For example:
        SAVED:/path/to/cover_shopify_senior-analytics-engineer_2026-06-22.docx
        SAVED:/path/to/cover_shopify_senior-analytics-engineer_2026-06-22.pdf
    """
    parser = argparse.ArgumentParser(description="Generate cover letter docx + pdf")
    parser.add_argument("--company",      required=True,  help="Normalised company slug (e.g. 'shopify')")
    parser.add_argument("--role",         required=True,  help="Normalised role slug (e.g. 'senior-analytics-engineer')")
    parser.add_argument("--content-file", required=True,  help="Path to a .py file defining a 'content' dict")
    parser.add_argument("--output-dir",   required=False, help="Destination directory (default: job-outputs/cover-letters/)")
    args = parser.parse_args()

    # Resolve the project root so all subsequent paths are absolute
    try:
        project_root = get_project_root()
    except RuntimeError as e:
        print(f"ERROR:{e}", file=sys.stderr)
        sys.exit(1)

    # Use the supplied output dir or fall back to the standard cover-letters folder
    output_dir = args.output_dir or os.path.join(project_root, "job-outputs", "cover-letters")
    os.makedirs(output_dir, exist_ok=True)

    # Build output filenames following the project naming convention
    today     = date.today().strftime("%Y-%m-%d")
    filename  = f"cover_{args.company}_{args.role}_{today}"
    docx_path = os.path.join(output_dir, filename + ".docx")
    pdf_path  = os.path.join(output_dir, filename + ".pdf")

    # Load the content dict from the caller-supplied Python file
    try:
        content = load_content(args.content_file)
    except Exception as e:
        print(f"ERROR:Failed to load content file: {e}", file=sys.stderr)
        sys.exit(1)

    # Create a blank document and apply project-standard margins before building
    doc = Document()
    set_margins(doc)

    try:
        build_cover_letter(doc, content, project_root)
    except Exception as e:
        print(f"ERROR:Document build failed: {e}", file=sys.stderr)
        sys.exit(1)

    doc.save(docx_path)

    # Re-open the saved file to confirm it is a valid docx (catches corruption)
    try:
        Document(docx_path)
    except Exception as e:
        print(f"ERROR:Generated docx failed validation: {e}", file=sys.stderr)
        sys.exit(1)

    # Convert to PDF; on failure, still report the docx as saved before exiting
    result_pdf = convert_to_pdf(docx_path, output_dir)
    if not result_pdf:
        print("ERROR:PDF conversion failed (soffice). docx saved but no pdf.", file=sys.stderr)
        # Emit the docx SAVED line so the skill can at least record the docx path
        print(f"SAVED:{docx_path}")
        sys.exit(1)

    # Emit SAVED: lines in docx-first order (the skill checks for both)
    print(f"SAVED:{docx_path}")
    if os.path.exists(pdf_path):
        print(f"SAVED:{pdf_path}")


if __name__ == "__main__":
    main()
