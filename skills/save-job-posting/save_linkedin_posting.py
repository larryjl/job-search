#!/usr/bin/env python3
"""
save_linkedin_posting.py — Save a LinkedIn job posting from extracted JD text.

This script accepts pre-extracted job description text (either from a file or
stdin) and formats it into a clean DOCX file inside ``job-outputs/postings/``.
It does NOT fetch the LinkedIn page itself; the calling agent is responsible for
extracting the JD text (e.g. via the Voyager API or a Chrome-based read) and
passing it here.

Why a separate script from ``save_job_posting.py``?  LinkedIn's raw HTML is
cluttered with navigation chrome, ads, and sign-in prompts.  Converting the URL
directly to PDF produces an unreadable artefact.  Instead the agent extracts
just the job description text and this script reformats it into a structured
DOCX — giving clean, searchable output without any LinkedIn UI noise.

Usage:
    python3 save_linkedin_posting.py \\
        --company "ShyftLabs" \\
        --title "Data Engineer" \\
        --url "https://www.linkedin.com/jobs/view/4385828168/" \\
        --jd-file /tmp/jd_text.txt

Or pipe JD text via stdin:
    echo "JD text..." | python3 save_linkedin_posting.py \\
        --company "ShyftLabs" --title "Data Engineer" \\
        --url "https://www.linkedin.com/jobs/view/4385828168/"

Output:
    Prints ``SAVED:/path/to/file`` and ``FILENAME:filename.docx`` on success.
"""

import argparse
import os
import re
import subprocess
import sys
from datetime import date

# Add the shared library directory to the import path so ``from project_paths
# import ...`` works regardless of the current working directory.
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../lib'))

from project_paths import get_project_root, get_postings_dir
from filename_builder import make_filename
from docx_helpers import set_metadata


def save_linkedin_posting(
    company: str,
    title: str,
    url: str,
    jd_text: str,
) -> tuple:
    """
    Format extracted LinkedIn JD text into a DOCX file and save it.

    The function attempts to reconstruct the posting's visual hierarchy by
    detecting section headers (short, capitalised lines) and bullet points
    (lines starting with ``-``, ``•``, or ``*``).  Everything else is treated
    as body text.

    The source URL is recorded as a plain paragraph near the top of the
    document so the reader can navigate back to the original posting.

    Args:
        company:  Company name — used in the heading and output filename.
        title:    Job title — used in the heading and output filename.
        url:      LinkedIn job URL (may be empty string if unknown).
        jd_text:  Full job description as plain text extracted from LinkedIn.

    Returns:
        A tuple of ``(output_path, filename)`` where ``output_path`` is the
        absolute path to the saved DOCX and ``filename`` is the basename.
    """
    # --- Auto-install python-docx if missing ---
    # We try the import first; if it fails we install it via pip and re-import.
    # This pattern avoids a hard dependency on the environment having python-docx
    # pre-installed while keeping the install invisible to the caller.
    try:
        from docx import Document
    except ImportError:
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx",
             "--break-system-packages", "-q"],
            check=True,
        )
        from docx import Document

    POSTINGS_DIR = get_postings_dir()

    doc = Document()

    # Stamp author/metadata into core properties for traceability.
    set_metadata(doc)

    # Level-1 heading identifies the role at a glance when the file is opened.
    doc.add_heading(f"{title} — {company}", level=1)

    # The source URL goes in its own paragraph — not in the heading — so that
    # the heading stays clean and the URL is easy to copy without selecting
    # extra text.  It is styled as Normal (not a heading) so it doesn't appear
    # in a table-of-contents outline.
    if url:
        meta = doc.add_paragraph()
        meta.add_run(f"Source: {url}").bold = False
        meta.style = doc.styles['Normal']

    # Blank spacer paragraph to visually separate the metadata block from the JD body.
    doc.add_paragraph()

    # --- Parse and render JD body ---
    # Split on double newlines to recover paragraph/section boundaries.
    # Single newlines within a section are preserved as separate lines.
    sections = jd_text.strip().split("\n\n")

    for section in sections:
        lines = section.strip().split("\n")
        if not lines:
            continue

        first_line = lines[0].strip()

        # Heuristic: treat a block as a section header when its first line is
        # short (< 80 chars), doesn't end with a period (ruling out body
        # sentences), and contains at least one uppercase letter (ruling out
        # all-lowercase fragments).  Single-line blocks that pass are rendered
        # as level-2 headings; multi-line blocks that pass are treated as body
        # text because the remaining lines are likely content, not a subheading.
        is_header = (
            len(first_line) < 80
            and not first_line.endswith('.')
            and len(lines) >= 1
            and any(c.isupper() for c in first_line)
        )

        if is_header and len(lines) == 1:
            # Lone short capitalised line → section heading (e.g. "About the Role")
            doc.add_heading(first_line, level=2)
        else:
            # Multi-line block or body paragraph — render line by line.
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Lines that start with common bullet markers are formatted as
                # list items using Word's built-in "List Bullet" style.
                # The first two characters (marker + space) are stripped so the
                # style's own bullet symbol is used instead of a doubled marker.
                if line.startswith(('- ', '• ', '* ')):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:])
                else:
                    doc.add_paragraph(line)

    # --- Build output path ---
    filename = make_filename(company, title, ext="docx")
    output_path = os.path.join(POSTINGS_DIR, filename)

    # If a file with this name already exists, append a numeric suffix (_2, _3,
    # …) to avoid silently overwriting a previously saved posting for the same
    # role.  We cap the search at suffix _9 to avoid an infinite loop.
    if os.path.exists(output_path):
        base = filename.rsplit('.', 1)[0]
        ext = filename.rsplit('.', 1)[1]
        for i in range(2, 10):
            candidate = os.path.join(POSTINGS_DIR, f"{base}_{i}.{ext}")
            if not os.path.exists(candidate):
                output_path = candidate
                filename = f"{base}_{i}.{ext}"
                break

    doc.save(output_path)
    return output_path, filename


def main() -> None:
    """Parse arguments, read JD text, call the save function, and print results."""
    parser = argparse.ArgumentParser(
        description="Save a LinkedIn job posting as a clean DOCX"
    )
    parser.add_argument("--company", required=True, help="Company name")
    parser.add_argument("--title", required=True, help="Job title")
    parser.add_argument("--url", default="", help="LinkedIn job URL")
    parser.add_argument(
        "--jd-file",
        help="Path to file containing JD text (or pipe via stdin)",
    )
    args = parser.parse_args()

    # Accept JD text from a file path or from stdin — whichever the caller provides.
    if args.jd_file:
        with open(args.jd_file) as f:
            jd_text = f.read()
    else:
        jd_text = sys.stdin.read()

    if not jd_text.strip():
        print("ERROR: No JD text provided", file=sys.stderr)
        sys.exit(1)

    output_path, filename = save_linkedin_posting(
        args.company, args.title, args.url, jd_text
    )

    # Machine-parseable output: the calling agent reads these two lines.
    print(f"SAVED:{output_path}")
    print(f"FILENAME:{filename}")


if __name__ == "__main__":
    main()
