#!/usr/bin/env python3
"""
Save a job posting as a PDF or DOCX file.

This script handles three input types:
- ``pasted``: raw text pasted by the agent — saved as a formatted DOCX.
- ``pdf``:    an uploaded PDF file — copied as-is to the postings directory.
- ``docx``:   an uploaded DOCX file — copied as-is to the postings directory.

URL-based inputs (fetching HTML from the web) are handled agent-side using
``url_to_pdf()`` from the shared lib; this script never makes network requests.
LinkedIn job URLs have their own dedicated script: ``save_linkedin_posting.py``.

Usage:
    python save_job_posting.py --input-type [pasted|pdf|docx] \\
        --company "Company Name" --title "Job Title" \\
        [--input-path /path/to/file | --input-text "..."]

Output:
    Prints ``SAVED:/path/to/saved/file`` and ``FILENAME:filename.ext`` on
    success, then exits 0.  Prints to stderr and exits 1 on failure.
"""

import sys
import os
import argparse
import shutil

# Make the shared library importable regardless of the working directory.
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../lib'))

from project_paths import get_postings_dir, get_project_root
from filename_builder import make_filename
from docx_helpers import set_metadata


def save_pasted_text(company: str, title: str, text: str, output_path: str) -> bool:
    """
    Convert raw pasted text into a formatted DOCX file and save it.

    Creates a simple document with a heading (``title — company``) followed
    by the job description text, one paragraph per line.  If ``python-docx``
    is not yet installed it is pip-installed automatically and the function
    retries itself once.

    Args:
        company:     Company name used in the document heading.
        title:       Job title used in the document heading.
        text:        Full job description text as a single string.
        output_path: Absolute path where the ``.docx`` file will be written.

    Returns:
        True if the file was written successfully, False on any error.
    """
    try:
        from docx import Document

        doc = Document()

        # Stamp author/metadata into the document's core properties so every
        # generated file is traceable back to the candidate.
        set_metadata(doc)

        # A level-1 heading makes the posting easy to identify when opened.
        doc.add_heading(f"{title} — {company}", level=1)

        # Add each line as its own paragraph to preserve the original
        # line-break structure of the pasted text (e.g. bullet points,
        # section headers).  We don't collapse lines because we don't know
        # which are headers and which are body text.
        for line in text.strip().split("\n"):
            doc.add_paragraph(line)

        doc.save(output_path)
        return True

    except ImportError:
        # python-docx is missing — install it silently, then retry.
        # We call the function recursively so the install-then-call pattern
        # stays self-contained without duplicating the creation logic.
        import subprocess
        subprocess.run(
            ["pip", "install", "python-docx", "--break-system-packages", "-q"],
            check=False,
        )
        return save_pasted_text(company, title, text, output_path)

    except Exception as e:
        print(f"Error saving pasted text: {e}", file=sys.stderr)
        return False


def copy_file(source_path: str, output_path: str) -> bool:
    """
    Copy an uploaded file to the postings directory under a normalized name.

    Uses ``shutil.copy`` (copy, not move) so the original uploaded file is
    preserved in its temporary location; the caller owns the lifecycle of the
    source file.

    Args:
        source_path: Absolute path to the source file (PDF or DOCX).
        output_path: Absolute destination path including filename.

    Returns:
        True if the copy succeeded, False on any error.
    """
    try:
        # shutil.copy preserves file content and permissions but does NOT
        # move the source — intentional, since the source may be a temp file
        # managed by the agent's upload machinery.
        shutil.copy(source_path, output_path)
        return True
    except Exception as e:
        print(f"Error copying file: {e}", file=sys.stderr)
        return False


def main() -> None:
    """
    Parse command-line arguments, build the output path, and dispatch to the
    appropriate save function based on ``--input-type``.
    """
    parser = argparse.ArgumentParser(
        description="Save a job posting as PDF or DOCX"
    )
    parser.add_argument(
        "--input-type",
        required=True,
        choices=["pasted", "pdf", "docx"],
        help=(
            "Type of input — 'url' inputs are handled agent-side; "
            "LinkedIn URLs use save_linkedin_posting.py"
        ),
    )
    parser.add_argument("--company", required=True, help="Company name")
    parser.add_argument("--title", required=True, help="Job title")
    parser.add_argument(
        "--input-path", help="Path to input file (for pdf/docx)"
    )
    parser.add_argument("--input-text", help="Pasted text (for pasted)")

    args = parser.parse_args()

    # Pasted text and uploaded DOCX files both produce a DOCX output.
    # Uploaded PDFs are copied verbatim, so they stay as PDF.
    ext = "docx" if args.input_type in ["pasted", "docx"] else "pdf"

    # Build the canonical filename (e.g. "acme_data-analyst_2026-06-22.docx")
    # and resolve the full destination path inside the postings directory.
    filename = make_filename(args.company, args.title, ext=ext)
    postings_dir = get_postings_dir()
    output_path = os.path.join(postings_dir, filename)

    # If a file with the same name already exists, append a numeric suffix
    # (_2, _3, …) rather than silently overwriting an existing posting.
    counter = 2
    base_path = output_path
    while os.path.exists(output_path):
        # Split off the extension, append the counter, then rejoin.
        # Example: "acme_role_2026-06-22.docx" → "acme_role_2026-06-22_2.docx"
        name_parts = base_path.rsplit(".", 1)
        output_path = f"{name_parts[0]}_{counter}.{name_parts[1]}"
        counter += 1
        filename = os.path.basename(output_path)

    # Dispatch to the correct handler for the input type.
    success = False
    if args.input_type == "pasted":
        success = save_pasted_text(args.company, args.title, args.input_text, output_path)
    elif args.input_type in ["pdf", "docx"]:
        success = copy_file(args.input_path, output_path)

    if success:
        # Machine-parseable output: the calling agent reads these two lines
        # to learn the final saved path and filename.
        print(f"SAVED:{output_path}")
        print(f"FILENAME:{filename}")
        sys.exit(0)
    else:
        print("ERROR: Failed to save job posting", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
