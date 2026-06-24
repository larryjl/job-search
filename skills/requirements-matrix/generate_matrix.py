#!/usr/bin/env python3
"""
Generate a requirements response matrix .docx from a structured content dict.

This script is used by the requirements-matrix skill to produce a formal
requirements response document for government or enterprise job applications
that require each mandatory and scored criterion to be addressed in writing.

The output is a single .docx file (no PDF conversion — the caller can convert
with soffice if needed).  Content is supplied via a .py file that defines a
top-level ``content`` dict, keeping the skill agent in control of the text
without needing to escape JSON special characters.

Usage:
    python3 skills/requirements-matrix/generate_matrix.py \\
        --organisation "government-of-alberta" \\
        --role "senior-data-architect" \\
        --content-file /tmp/matrix_content.py \\
        --output-dir job-outputs/matrices/

Content file format (/tmp/matrix_content.py):
    content = {
        "organisation": "Government of Alberta",   # display name (title case)
        "role": "Senior Data Architect",            # display name
        "mandatory": [
            {
                "label": "M1",
                "wording": "Minimum 7 years experience with enterprise data platforms...",
                "verdict": "Met.",
                "summary": "7+ years across three roles.",   # one-sentence summary after verdict
                "bullets": [
                    "ACME Corp, Senior Analyst (2019–2025): Led design of Snowflake data warehouse...",
                    "Previous Co, Data Analyst (2016–2019): Built ETL pipelines supporting...",
                ],
            },
        ],
        "scored": [
            {
                "label": "S1",
                "wording": "Experience with cloud-native data platforms (AWS, Azure, GCP)...",
                "verdict": "Met.",
                "summary": "AWS and Azure experience across multiple production deployments.",
                "bullets": [...],
            },
        ],
    }

Output:
    Prints SAVED:/path/to/file.docx on success.
    Prints ERROR:... and exits with code 1 on failure.
"""

import argparse
import importlib.util
import os
import sys
from datetime import date

# Store the directory containing this script so relative lib paths are stable
# regardless of whatever working directory the caller uses when invoking the script.
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def _setup_lib_path() -> None:
    """
    Prepend the shared skills/lib directory to sys.path if not already present.

    All skill scripts that use the shared library (project_paths, docx_helpers,
    etc.) call this function at import time so the shared modules are importable
    without needing to install them as a package.
    """
    lib_path = os.path.normpath(os.path.join(SCRIPT_DIR, '../lib'))
    if lib_path not in sys.path:
        sys.path.insert(0, lib_path)


# Set up the path before importing shared modules.
_setup_lib_path()

from project_paths import get_project_root   # noqa: E402 (import after path setup)
from docx_helpers import (                   # noqa: E402
    NAVY, BLACK,
    set_spacing, add_run, set_margins, keep_with_next,
    add_rule, add_bullet, set_metadata,
)
from docx import Document                    # noqa: E402


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_content(content_file: str) -> dict:
    """
    Dynamically import a Python file and return its ``content`` dict.

    The content file is a plain .py file that defines a single top-level
    variable called ``content``.  Using ``importlib`` to load it (rather than
    ``exec()`` or a regular ``import``) lets us load an arbitrary file path
    as a module without adding its directory to sys.path or risking name
    collisions with real modules already in the interpreter.

    Args:
        content_file: Absolute or relative path to the .py content file.

    Returns:
        The ``content`` dict defined inside the file.

    Raises:
        ValueError: If the loaded file does not define a variable named ``content``.
    """
    # ``spec_from_file_location`` creates a module spec telling Python how to
    # load the file — essentially "treat this path as a module called matrix_content".
    spec = importlib.util.spec_from_file_location("matrix_content", content_file)

    # ``module_from_spec`` allocates an empty module object from that spec.
    module = importlib.util.module_from_spec(spec)

    # ``exec_module`` actually runs the file, populating ``module`` with all
    # top-level variables (including ``content``).
    spec.loader.exec_module(module)

    if not hasattr(module, 'content'):
        raise ValueError(
            f"content_file must define a variable named 'content': {content_file}"
        )
    return module.content


# ---------------------------------------------------------------------------
# Document builder helpers
# ---------------------------------------------------------------------------

def add_section_heading(doc: Document, text: str) -> object:
    """
    Add an ALL-CAPS section heading in 12pt bold navy.

    This is a matrix-specific heading helper — it uses slightly different
    spacing from the resume version in docx_helpers.add_section_heading
    (8pt before vs 10pt) and does not add a left border, keeping the matrix
    document style clean and formal.

    Args:
        doc:  The python-docx Document to append the heading to.
        text: Heading text; will be uppercased automatically.

    Returns:
        The newly added paragraph.
    """
    p = doc.add_paragraph()
    set_spacing(p, before=8, after=4)

    # keep_with_next ensures the heading is never stranded at the bottom of a
    # page without at least the first line of content that follows it.
    keep_with_next(p)

    add_run(p, text.upper(), bold=True, size=12, color=NAVY)
    return p


def add_requirement(doc: Document, req: dict) -> None:
    """
    Render one requirement block: label + wording, verdict line, and bullets.

    Each requirement in the content dict follows a standard three-part layout:
      1. Bold label + requirement wording (e.g. "M1 — Minimum 7 years experience...")
      2. Bold verdict + one-sentence summary (e.g. "Met. 7+ years across three roles.")
      3. Bullet points providing specific evidence (one bullet per relevant role)

    Args:
        doc: The python-docx Document to append the requirement block to.
        req: A dict with keys: "label", "wording", "verdict", "summary", "bullets".
             All keys except "label" and "wording" are optional.
    """
    # --- Label + wording line (bold, 11pt) ---
    p = doc.add_paragraph()
    set_spacing(p, before=8, after=2)
    keep_with_next(p)  # keep label attached to the verdict line below

    # Combine label and wording into a single bold line so the reader can
    # quickly scan which criterion is being addressed.
    label_text = f"{req['label']} — {req['wording']}"
    add_run(p, label_text, bold=True, size=11)

    # --- Verdict + summary line (bold, 11pt) ---
    verdict_text = req.get("verdict", "")
    summary_text = req.get("summary", "")

    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2)
    keep_with_next(p)  # keep verdict attached to the first bullet below

    # Concatenate verdict and summary into one paragraph (e.g. "Met. 7+ years...").
    # If there is no summary, only the verdict is shown.
    full_verdict = verdict_text
    if summary_text:
        full_verdict += f" {summary_text}"
    add_run(p, full_verdict, bold=True, size=11)

    # --- Evidence bullets ---
    # Each bullet is a specific experience item from the candidate's background.
    for bullet_text in req.get("bullets", []):
        add_bullet(doc, bullet_text)

    # Spacer paragraph adds visual breathing room between requirements.
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6)


def build_matrix(doc: Document, content: dict, project_root: str) -> None:
    """
    Populate a blank python-docx Document with requirements matrix content.

    Renders sections in this order:
        1. Title block — candidate name, organisation/role, document type, date
        2. Horizontal rule
        3. Mandatory requirements (one block per item in content["mandatory"])
        4. Horizontal rule
        5. Scored requirements (one block per item in content["scored"])

    Sections with empty or absent lists are silently skipped.

    Args:
        doc:          A freshly created Document with margins already set.
        content:      Structured content dict (see module docstring for schema).
        project_root: Absolute path to the job-search project root; passed to
                      set_metadata to read the candidate name from master-resume.md.
    """
    # Write author/title into DOCX core properties and return the candidate name.
    candidate_name = set_metadata(doc, project_root)

    organisation = content.get("organisation", "")
    role         = content.get("role", "")
    today_str    = date.today().strftime("%B %-d, %Y")  # e.g. "June 22, 2026"

    # --- Title block ---
    # Candidate name: large, bold, navy — mirrors the resume header style.
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2)
    add_run(p, candidate_name, bold=True, size=16, color=NAVY)

    # Organisation and role on one line, slightly smaller than the name.
    if organisation and role:
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=8)
        add_run(p, f"{organisation} — {role}", size=12)

    # Document type label — tells the reader what this document is at a glance.
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2)
    add_run(p, "Requirement Response Matrix", size=12)

    # Date the document was generated.
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=8)
    add_run(p, today_str, size=12)

    # Horizontal rule separates the title block from the body.
    add_rule(doc)

    # --- Mandatory requirements ---
    mandatory = content.get("mandatory", [])
    if mandatory:
        add_section_heading(doc, "Mandatory Requirements")
        for req in mandatory:
            add_requirement(doc, req)
        add_rule(doc)  # rule after mandatory block, before scored

    # --- Scored requirements ---
    scored = content.get("scored", [])
    if scored:
        add_section_heading(doc, "Scored Requirements")
        for req in scored:
            add_requirement(doc, req)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    """
    Entry point: parse arguments, build the matrix document, save and validate.

    Exit codes:
        0 — docx saved and validated successfully
        1 — any error (content load failure, build failure, validation failure)

    Output on success:
        SAVED:/absolute/path/to/matrix_<organisation>_<role>_<date>.docx
    The SAVED: prefix is a machine-readable protocol parsed by the calling skill.
    """
    parser = argparse.ArgumentParser(description="Generate requirements matrix docx")
    parser.add_argument(
        "--organisation",
        required=True,
        help="Normalised organisation slug (e.g. 'government-of-alberta')",
    )
    parser.add_argument(
        "--role",
        required=True,
        help="Normalised role slug (e.g. 'senior-data-architect')",
    )
    parser.add_argument(
        "--content-file",
        required=True,
        help="Path to a .py file defining a 'content' dict",
    )
    parser.add_argument(
        "--output-dir",
        required=False,
        help="Destination directory (default: job-outputs/matrices/)",
    )
    args = parser.parse_args()

    # Resolve the project root so all subsequent paths are absolute.
    try:
        project_root = get_project_root()
    except RuntimeError as e:
        print(f"ERROR:{e}", file=sys.stderr)
        sys.exit(1)

    # Use the supplied output dir or fall back to the standard matrices folder.
    output_dir = args.output_dir or os.path.join(project_root, "job-outputs", "matrices")
    os.makedirs(output_dir, exist_ok=True)  # create the directory if it doesn't exist

    # Build the output filename following the project naming convention.
    today     = date.today().strftime("%Y-%m-%d")
    filename  = f"matrix_{args.organisation}_{args.role}_{today}"
    docx_path = os.path.join(output_dir, filename + ".docx")

    # Load the structured content dict from the caller-supplied Python file.
    try:
        content = load_content(args.content_file)
    except Exception as e:
        print(f"ERROR:Failed to load content file: {e}", file=sys.stderr)
        sys.exit(1)

    # Create a blank document and apply project-standard page margins.
    doc = Document()
    set_margins(doc)

    # Build and populate the document.
    try:
        build_matrix(doc, content, project_root)
    except Exception as e:
        print(f"ERROR:Document build failed: {e}", file=sys.stderr)
        sys.exit(1)

    doc.save(docx_path)

    # Re-open the saved file to confirm it is a valid docx — python-docx can
    # occasionally produce invalid XML for certain content; this catches it early.
    try:
        Document(docx_path)
    except Exception as e:
        print(f"ERROR:Generated docx failed validation: {e}", file=sys.stderr)
        sys.exit(1)

    # Machine-readable output: the calling skill parses this line to get the path.
    print(f"SAVED:{docx_path}")


if __name__ == "__main__":
    main()
