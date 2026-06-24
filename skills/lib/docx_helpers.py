"""
Shared python-docx helpers for all document-generating skills.

This is the single source of truth for fonts, colours, margins, and layout
helpers. Skills import from here — never copy-paste these definitions inline.

Exported symbols
----------------
Constants:
    NAVY        RGBColor — #1F3765; used for headings and accent borders
    BLACK       RGBColor — #1A1A1A; used for body text, dates, contact info
    FONT        str      — "Arial"; used for body text, bullets, and headings
    FONT_NAME   str      — "Georgia"; used exclusively for the candidate name
    MARGIN_IN   float    — 0.85; page margin in inches applied to all sides

Core helpers:
    set_spacing(paragraph, before, after, line)  — paragraph spacing in pt
    add_run(paragraph, text, bold, size, color)  — styled text run
    set_margins(doc, inches)                     — apply equal page margins
    keep_with_next(paragraph)                    — suppress page break after

Resume-specific helpers:
    add_company_date_row(doc, company, date_range) — left company / right date
    add_prof_dev_item(doc, credential, description) — bold credential + desc
    add_section_heading(doc, text)               — ALL-CAPS heading + left border

Requirements-matrix helpers:
    add_rule(doc, color, size)  — horizontal rule via paragraph bottom border
    add_bullet(doc, text, indent) — hanging-indent bullet point

Metadata:
    set_metadata(doc, project_root) — set author/last_modified_by from resume

Usage:
    import sys, os
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../lib'))
    from docx_helpers import (
        NAVY, BLACK, FONT, FONT_NAME, MARGIN_IN,
        set_spacing, add_run, set_margins, keep_with_next,
        add_section_heading, add_company_date_row, add_rule, add_bullet,
        add_prof_dev_item, set_metadata,
    )
"""

import os
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

NAVY  = RGBColor(0x1F, 0x37, 0x65)   # #1F3765 — headings, name, accent
BLACK = RGBColor(0x1A, 0x1A, 0x1A)   # #1A1A1A — body text, dates, contact
FONT       = "Arial"    # body text, bullets, section headings
FONT_NAME  = "Georgia"  # candidate name only (header)
MARGIN_IN  = 0.85       # page margins in inches


# ---------------------------------------------------------------------------
# Core helpers
# ---------------------------------------------------------------------------

def set_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph space_before, space_after, and optionally line_spacing (all in pt)."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def add_run(paragraph, text, bold=False, size=11, color=None):
    """
    Add a run to paragraph with standard font settings.

    Args:
        paragraph: docx Paragraph object
        text (str): Run text
        bold (bool): Bold weight
        size (int|float): Font size in pt
        color: RGBColor — defaults to BLACK

    Returns:
        Run object
    """
    if color is None:
        color = BLACK
    run = paragraph.add_run(text)
    run.bold           = bold
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.font.name      = FONT
    return run


def set_margins(doc, inches=None):
    """Set equal margins on all sides (default MARGIN_IN = 0.85 inch)."""
    if inches is None:
        inches = MARGIN_IN
    for section in doc.sections:
        section.top_margin    = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin   = Inches(inches)
        section.right_margin  = Inches(inches)


def keep_with_next(paragraph):
    """Prevent a page break between this paragraph and the next."""
    paragraph.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# Resume-specific helpers
# ---------------------------------------------------------------------------

def add_company_date_row(doc, company, date_range):
    """
    Add a paragraph with company name left-aligned and date range flush right.

    The flush-right effect is achieved with a single right-aligned tab stop
    positioned at 6.8 inches from the left edge — calculated to sit at the
    right margin of a US Letter page (8.5" wide minus 2 × 0.85" margins = 6.8").

    Tab-stop positions in OOXML are expressed in twips (twentieths of a point,
    i.e. 1/1440 of an inch), so 6.8 inches × 1440 = 9792 twips.

    Args:
        doc: Document object
        company (str): Company name (bold, BLACK)
        date_range (str): Date range string (regular, BLACK)

    Returns:
        Paragraph object
    """
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)

    # --- Build a right-aligned tab stop via raw OOXML ---
    # OOXML paragraph properties live in the <w:pPr> element; get or create it.
    pPr  = p._p.get_or_add_pPr()

    # <w:tabs> is the container for one or more tab-stop definitions.
    tabs = OxmlElement('w:tabs')

    # <w:tab> defines a single tab stop:
    #   w:val="right"  — text to the left of the tab is right-aligned at this stop
    #   w:pos          — position in twips from the left page edge
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    # 6.8 inches × 1440 twips/inch = 9792 twips
    # This places the tab stop flush with the right margin on a US Letter page
    # with 0.85" margins on each side (8.5" - 0.85" - 0.85" = 6.8" usable width).
    tab.set(qn('w:pos'), str(int(6.8 * 1440)))

    tabs.append(tab)
    pPr.append(tabs)

    # Company name: bold, left-aligned (appears before the tab character)
    run_co = p.add_run(company)
    run_co.bold           = True
    run_co.font.color.rgb = BLACK
    run_co.font.name      = FONT

    # The tab character jumps to the right tab stop, pushing the date to the right
    p.add_run('\t')

    # Date range: plain weight, right-aligned at the tab stop
    run_date = p.add_run(date_range)
    run_date.bold           = False
    run_date.font.color.rgb = BLACK
    run_date.font.name      = FONT

    return p


def add_prof_dev_item(doc, credential, description=None):
    """
    Add a Professional Development entry as two separate paragraphs.

    The credential is bold; the description (if any) is a plain paragraph
    on the next line. Never concatenate them inline.

    Args:
        doc: Document object
        credential (str): e.g. "dbt Analytics Engineering | 2024"
        description (str|None): Optional plain-text description line

    Returns:
        tuple: (credential_paragraph, description_paragraph or None)
    """
    p_cred = doc.add_paragraph()
    # Collapse bottom spacing when a description follows — the description
    # paragraph provides its own 6pt bottom gap instead.
    set_spacing(p_cred, before=0, after=0 if description else 6)
    add_run(p_cred, credential, bold=True)

    p_desc = None
    if description:
        p_desc = doc.add_paragraph()
        set_spacing(p_desc, before=0, after=6)
        add_run(p_desc, description)

    return p_cred, p_desc


# ---------------------------------------------------------------------------
# Resume section heading
# ---------------------------------------------------------------------------

def add_section_heading(doc, text):
    """
    Add an ALL-CAPS section heading (12pt bold NAVY) with a 2pt navy left border.

    The left border is implemented as an OOXML paragraph border (<w:pBdr>) rather
    than a drawing element, so it travels with the text and is ATS-safe (parsers
    see the text, not a floating shape).

    OOXML border attributes used:
        w:pBdr  — paragraph border container; holds directional border children
        w:left  — specifies the left-side border of the paragraph
        w:val   — border style (e.g. "single" = solid line)
        w:sz    — border thickness in eighths of a point (16 eighths = 2pt)
        w:space — gap in points between the border line and the paragraph text
        w:color — border colour as a 6-digit hex string (no '#' prefix)

    Args:
        doc: Document object
        text (str): Heading text (will be uppercased)

    Returns:
        Paragraph object
    """
    p = doc.add_paragraph()
    set_spacing(p, before=10, after=3)
    # Keep heading visually attached to the content that follows it
    keep_with_next(p)

    # --- Attach a left border via raw OOXML ---
    # Retrieve (or create) the paragraph properties element <w:pPr>
    pPr  = p._p.get_or_add_pPr()

    # <w:pBdr> is the paragraph border container element
    pBdr = OxmlElement('w:pBdr')

    # <w:left> sets only the left-side border; other sides remain unset
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')   # solid line style
    left.set(qn('w:sz'),    '16')       # thickness: 16 eighths-of-a-point = 2pt
    left.set(qn('w:space'), '8')        # 8pt gap between border line and text
    left.set(qn('w:color'), '1F3765')   # NAVY — matches the NAVY constant above

    pBdr.append(left)
    pPr.append(pBdr)

    add_run(p, text.upper(), bold=True, size=12, color=NAVY)
    return p


# ---------------------------------------------------------------------------
# Requirements-matrix helpers
# ---------------------------------------------------------------------------

def add_rule(doc, color="2E5496", size=12):
    """
    Draw a horizontal rule using a bottom border on a blank paragraph.

    Rather than inserting a drawing or a table, this attaches a <w:bottom>
    paragraph border to an otherwise-empty paragraph. The result looks like
    a horizontal rule and is fully ATS-safe — no shapes or images involved.

    w:bottom is the OOXML border element for the bottom edge of a paragraph,
    analogous to w:left used in add_section_heading but spanning the full
    text width instead of sitting on the side.

    Args:
        doc: Document object
        color (str): Hex color string without '#' (default: "2E5496")
        size (int): Border thickness in eighths of a point (default: 12 = 1.5pt)

    Returns:
        Paragraph object
    """
    p   = doc.add_paragraph()
    set_spacing(p, before=4, after=4)

    # Build the OOXML border structure on the blank paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr   = OxmlElement('w:pBdr')

    # <w:bottom> renders as a horizontal line across the full text width
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(size))
    bottom.set(qn('w:space'), '1')       # 1pt gap below text (paragraph is empty)
    bottom.set(qn('w:color'), color)

    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, indent=0.25):
    """
    Add a hanging-indent bullet point paragraph.

    Args:
        doc: Document object
        text (str): Bullet text (bullet char prepended automatically)
        indent (float): Left indent in inches (default 0.25)

    Returns:
        Paragraph object
    """
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2)
    # left_indent pushes the whole paragraph right; first_line_indent pulls
    # the first line (with the bullet) back to the left margin, creating the
    # classic hanging-indent look where wrapped lines align under the text,
    # not under the bullet character.
    p.paragraph_format.left_indent       = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-indent)
    add_run(p, "• " + text)
    return p


# ---------------------------------------------------------------------------
# Metadata
# ---------------------------------------------------------------------------

def set_metadata(doc, project_root=None):
    """
    Set standard docx core properties from profile/master-resume.md.

    Reads the first line of master-resume.md (strips leading '# ') to get
    the candidate name, then sets author, last_modified_by, and clears comments.

    Both `author` and `last_modified_by` are set to the same name because
    Word and LibreOffice track them independently: `author` records who
    originally created the file, while `last_modified_by` records who saved
    it most recently. Setting both prevents the generated file from showing
    the build machine's OS username (e.g. "root" or "runner") as the last
    editor when the document is opened on a different system.

    Args:
        doc: Document object
        project_root (str|None): Override project root path. If None, resolved
            via project_paths.get_project_root().

    Returns:
        str: Candidate name that was set
    """
    if project_root is None:
        # Support both package-relative imports (when used as part of the lib
        # package) and sys.path-based imports (when called from a skill script).
        try:
            from .project_paths import get_project_root
        except ImportError:
            from project_paths import get_project_root
        project_root = get_project_root()

    resume_path = os.path.join(project_root, "profile", "master-resume.md")
    with open(resume_path) as f:
        # First line of the resume is the candidate's name in Markdown H1 format:
        # "# Lawrence Lee" → strip the leading "# " → "Lawrence Lee"
        candidate_name = f.readline().lstrip("# ").strip() or "Unknown"

    doc.core_properties.author           = candidate_name
    doc.core_properties.last_modified_by = candidate_name
    doc.core_properties.comments         = ""  # clear any auto-generated comments
    return candidate_name
